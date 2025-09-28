import streamlit as st
import openpyxl
from openpyxl.utils.cell import get_column_letter
import pandas as pd
import io
import os
import re
import zipfile
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Set Streamlit page configuration
st.set_page_config(
    page_title="Excel Export to Word Form Generator",
    layout="wide",
    initial_sidebar_state="expanded"
)


# --- Utility Functions for Streamlit ---

def st_display_warning(title, message):
    """Replaces tkinter.messagebox.showwarning with Streamlit equivalent."""
    st.warning(f"**{title}**\n\n{message}")


def st_display_error(title, message):
    """Replaces tkinter.messagebox.showerror with Streamlit equivalent."""
    st.error(f"**{title}**\n\n{message}")


def add_cell_border(cell, color_rgb=(0x00, 0x00, 0x00), size_pt=4):
    """Helper to add borders to a docx table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Define border attributes
    borders = ['top', 'left', 'bottom', 'right']
    for border in borders:
        # Create the border element
        tag = qn(
            f'w:{border}')
        e = OxmlElement(tag)
        e.set(qn('w:val'), 'single')  # Border style
        e.set(qn('w:sz'), str(size_pt))  # Size in 1/8 of a point
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), f'{color_rgb[0]:02X}{color_rgb[1]:02X}{color_rgb[2]:02X}')  # Color in hex
        tcPr.append(e)


# --- REVISED: Aggressive Text cleaning for DOCX insertion ---
def safe_text_for_docx(value):
    """
    Converts a value to string and strips characters that can cause XML parsing errors
    in python-docx, such as control characters and stray XML/HTML fragments.
    """
    if value is None:
        return ""

    s = str(value)

    # 1. Remove non-ASCII/non-printable control characters using encode/decode
    s = s.encode('ascii', 'ignore').decode('ascii')

    # 2. Clean up common XML/HTML-like fragments that cause the reported error
    # Target the known fragments like '{http'
    s = re.sub(r'\{http.*?\}', '', s, flags=re.IGNORECASE)
    s = re.sub(r'<[^>]+>', '', s)  # Remove stray HTML/XML tags
    s = s.replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&') # Decode basic HTML entities

    # 3. Explicitly remove non-breaking spaces (can also cause issues) and trim
    return s.replace('\xa0', ' ').strip()


# --- Excel Pre-processing Function (PHASE 1) ---
def process_excel_st(uploaded_file, sheet_name="Worksheet", start_row=1):
    """
    Takes the uploaded form/fields export and processes it in memory (filling blanks, deleting rows).
    """
    try:
        # Load workbook from the uploaded file's bytes buffer
        workbook_data = uploaded_file.read()
        wb = openpyxl.load_workbook(io.BytesIO(workbook_data))

        # Check if the required sheet exists
        if sheet_name not in wb.sheetnames:
            st_display_error("Error",
                             f"Sheet '{sheet_name}' not found in the uploaded file. Available sheets: {', '.join(wb.sheetnames)}")
            return None, None

        sheet = wb[sheet_name]

        st.info("Starting Excel Pre-processing...")

        end_row_for_processing = sheet.max_row
        end_row_for_deletion = sheet.max_row

        # Step 1: Find the "Bird species" row and determine the actual end_row
        bird_species_found_at_row = None
        for row_index in range(start_row, sheet.max_row + 1):
            # Check cell 7 (G) for "Bird species"
            cell_value_col7 = sheet.cell(row=row_index, column=7).value
            if isinstance(cell_value_col7, str) and cell_value_col7.strip() == "Bird species":
                bird_species_found_at_row = row_index
                break

        if bird_species_found_at_row is not None:
            # Check cell 8 (H) for "Position ID"
            position_id_value = sheet.cell(row=bird_species_found_at_row, column=8).value
            if position_id_value is None or str(position_id_value).strip() == "":
                st.write(
                    f"ℹ️ **'Bird species' found** with a BLANK 'Position ID'. Rows after row {bird_species_found_at_row - 1} **WILL BE removed**.")
                end_row_for_processing = bird_species_found_at_row - 1
                end_row_for_deletion = bird_species_found_at_row
            else:
                st.write(
                    f"ℹ️ **'Bird species' found** with a valid 'Position ID'. Rows after this point **WILL NOT be removed**.")
        else:
            st.write("ℹ️ **'Bird species' not found**. All rows will be processed for filling blanks.")

        if end_row_for_processing < start_row:
            st_display_warning("Excel Processing Warning",
                               f"Calculated end row for processing ({end_row_for_processing}) is less than start_row ({start_row}). No filling operations will be performed.")
            # Still save the workbook in case it had header content
            output_buffer = io.BytesIO()
            wb.save(output_buffer)
            # The second return value (valid_values) is not used if processing is skipped, but we return an empty set for consistency.
            return output_buffer.getvalue(), set()

        if end_row_for_processing >= start_row:
            # Fill blank cells in columns 1 to 10 (excluding headers, hence start_row + 1)
            st.write("Filling blanks in columns 1-10...")
            for col_index in range(1, 11):
                col_letter = get_column_letter(col_index)
                for row_index in range(start_row + 1, end_row_for_processing + 1):
                    current_cell = sheet[f"{col_letter}{row_index}"]
                    if current_cell.value is None or (
                            isinstance(current_cell.value, str) and current_cell.value.strip() == ""):
                        above_cell = sheet[f"{col_letter}{row_index - 1}"]
                        current_cell.value = above_cell.value

            # Special fill for Column 4 based on Column 3
            st.write("Applying special fill for Column 4...")
            for row_index in range(start_row + 1, end_row_for_processing + 1):
                col_c_value = sheet.cell(row=row_index, column=3).value
                col_d_cell = sheet.cell(row=row_index, column=4)

                if (col_c_value is not None and str(col_c_value).strip() != "") and \
                        (col_d_cell.value is None or str(col_d_cell.value).strip() == ""):
                    col_d_cell.value = "(empty subsection)"

        # Define valid values for column 6 for conditional filling of columns 14-15
        # NOTE: This list of Field IDs (valid_values) is still needed for conditional filling in the Excel processing.
        valid_values = {10, 11, 12, 14, 15, 17, 18, 19, 28, 29, 30, 94, 103, 104, 107, 108, 109, 117,
                        119, 120, 122, 128, 130, 131, 142, 160, 168, 170, 206, 208, 209, 212, 214, 216,
                        217, 220, 221, 222, 223, 224, 225, 226, 227, 228, 284, 292, 293, 311, 314, 315,
                        316, 320, 325, 326, 371, 431, 980, 995, 1391, 1401, 1404, 1672, 1673, 1759, 2329,
                        2440, 2472, 2473, 2485, 2486, 2500, 2502, 3282, 3708, 3709, 3875, 4240, 4243,
                        4244, 4245, 5894, 5901, 5940, 7349, 8829, 8830, 9114, 9115, 9116, 9117, 9118, 9119,
                        9120, 9121, 9122, 9126, 9127, 9128, 9129, 9130, 9135, 9136, 9137, 9138, 9139, 9140,
                        9141, 9142, 9143, 9144, 9145, 9150, 9151, 9152, 9153, 9154, 9155, 9156, 9157, 9158,
                        9179, 9180, 9181, 9189, 9191, 9192, 9201, 9202, 9203, 9204, 9205, 9206, 9207, 9208,
                        9276, 9277, 9278, 9279, 9280, 9281, 9282, 9298, 9299, 9333, 9529, 9540, 9543, 9544,
                        9552, 9557, 9565, 10614, 10618, 12739, 12960, 15463, 23157, 23159}

        if end_row_for_processing >= start_row:
            # Fill blank cells in columns 14 to 15 based on conditions
            st.write("Applying conditional fill for columns 14-15...")
            for row_index in range(start_row + 1, end_row_for_processing + 1):
                col_10_value = sheet.cell(row=row_index, column=10).value
                try:
                    # Column 6 (F) is 'Field ID', which is sometimes numeric
                    col_6_value = int(sheet.cell(row=row_index, column=6).value)
                except (ValueError, TypeError):
                    col_6_value = None

                # Condition: Field Type is "Dropdown select" (Col 10) AND Field ID (Col 6) is in the valid list
                if col_10_value == "Dropdown select" and col_6_value in valid_values:
                    for col_index in range(14, 16):  # Checks columns N (14) and O (15)
                        col_letter = get_column_letter(col_index)
                        current_cell = sheet[f"{col_letter}{row_index}"]
                        if current_cell.value is None:
                            above_cell = sheet[f"{col_letter}{row_index - 1}"]
                            current_cell.value = above_cell.value

        # Delete rows if "Bird species" criteria was met
        if sheet.max_row >= end_row_for_deletion:
            amount_to_delete = sheet.max_row - end_row_for_deletion + 1
            if amount_to_delete > 0:
                st.write(f"🗑️ Removed {amount_to_delete} rows starting from row {end_row_for_deletion}.")
                sheet.delete_rows(end_row_for_deletion, amount_to_delete)

        # Save the processed workbook to a BytesIO object
        output_buffer = io.BytesIO()
        wb.save(output_buffer)

        return output_buffer.getvalue(), valid_values

    except KeyError:
        # This is now handled with an explicit check above, but keeping the catch-all
        st_display_error("Error", f"Sheet '{sheet_name}' not found in the uploaded file.")
        return None, None
    except Exception as e:
        st_display_error("An Error Occurred", f"An error occurred during Excel processing: {e}")
        return None, None


# --- Word Form Generation Function (PHASE 3: READS Session State) ---
def create_forms_from_excel_st(excel_bytes, folder_name):
    """
    Generates Word documents based on the processed Excel, reading large option choices
    from st.session_state.
    """
    try:
        # Read the Excel data from bytes into a DataFrame
        df = pd.read_excel(io.BytesIO(excel_bytes))

        st.info("Starting Word Form Generation...")

        # Column standardization (remains the same)
        df["Form Description"] = df["Form Description"].astype(str).fillna("")
        df["Section"] = df["Section"].astype(str).fillna("")
        df["Subsection Header"] = df["Subsection Header"].astype(str).replace("nan", "n/a").fillna("")
        df["Field ID"] = df["Field ID"].astype(str).fillna("")
        df["Field Description"] = df["Field Description"].astype(str).fillna("")
        df["Form ID"] = df["Form ID"].astype(str).fillna("")
        df["Field Type"] = df["Field Type"].astype(str).fillna("")
        df["Option"] = df["Option"].astype(str).replace("nan", "n/a").fillna("")
        df["Mandatory"] = df["Mandatory"].astype(str).fillna("")
        df["Eccairs Value ID"] = df["Eccairs Value ID"].astype(str).replace("nan", "n/a").fillna("")
        df["Eccairs Value"] = df["Eccairs Value"].astype(str).replace("nan", "n/a").fillna("")

        generated_files = []
        today_date_str = date.today().strftime("%Y%m%d")

        for (form_desc, form_id), form_group in df.groupby(["Form Description", "Form ID"]):
            document = Document()
            document.add_heading(f'Form: {form_desc} [{form_id}]', level=0)

            # Sort sections by their minimum Position ID
            unique_sections = form_group.groupby(form_group['Section'].astype(str))[
                'Position ID'].min().sort_values().index

            for section in unique_sections:
                section_group = form_group[form_group['Section'].astype(str) == section]
                document.add_heading(section, level=1)

                # Sort subsections by their minimum Position ID
                unique_subsections = section_group.groupby(section_group['Subsection Header'].astype(str))[
                    'Position ID'].min().sort_values().index

                for subsection_header in unique_subsections:
                    if str(subsection_header).strip().lower() not in ("n/a", ""):
                        document.add_heading(subsection_header, level=2)
                        subsection_group = section_group[
                            section_group['Subsection Header'].astype(str) == subsection_header]
                    else:
                        subsection_group = section_group[
                            (section_group['Subsection Header'].astype(str).str.strip().str.lower() == "n/a") |
                            (section_group['Subsection Header'].astype(str).str.strip() == "")
                            ]
                        if subsection_group.empty:
                            continue

                    processed_field_ids_for_current_subsection = set()

                    # Get unique fields in the current subsection, sorted by Position ID
                    unique_fields_in_subsection = subsection_group.sort_values(
                        by="Position ID").drop_duplicates(subset=["Field ID"])

                    for index, row in unique_fields_in_subsection.iterrows():
                        field_id = str(row["Field ID"])  # Ensure Field ID is treated as a string for comparison
                        field_desc = row["Field Description"]
                        field_type = row["Field Type"]
                        is_mandatory = str(row.get("Mandatory", "")).strip().upper() == "T"

                        if field_id in processed_field_ids_for_current_subsection:
                            continue

                        # Apply safe_text_for_docx here for the display text as well, just in case
                        display_field_desc = safe_text_for_docx(field_desc)
                        display_field_desc = f"{display_field_desc}*" if is_mandatory else display_field_desc
                        safe_field_id = safe_text_for_docx(field_id)

                        if field_type == "Dropdown select":
                            p = document.add_paragraph()
                            p.add_run(f'{display_field_desc}: ').bold = True
                            p.add_run(f'[{field_type}; {safe_field_id}]')

                            # Fetch options unique to this Field ID
                            options_for_field = df[
                                (df["Field ID"].astype(str) == field_id) &
                                (df["Option"].astype(str).str.lower() != "n/a") &
                                (df["Option"].notna())
                                ]["Option"].unique().tolist()

                            if options_for_field:
                                if len(options_for_field) > 50:
                                    # Read the stored choice made in Phase 2
                                    choice = st.session_state.get(f'config_choice_{field_id}', False)

                                    if choice:
                                        p_options_label = document.add_paragraph("Options:", style='Normal')
                                        p_options_label.paragraph_format.left_indent = Inches(0.5)

                                        for option in options_for_field:
                                            # Use safe_text_for_docx here for high-volume content
                                            safe_option = safe_text_for_docx(option)
                                            p_option_bullet = document.add_paragraph(safe_option, style='List Bullet')
                                            p_option_bullet.paragraph_format.left_indent = Inches(1.0)
                                    else:
                                        p_options_label = document.add_paragraph("Options:", style='Normal')
                                        p_options_label.paragraph_format.left_indent = Inches(0.5)
                                        p_grouped_options = document.add_paragraph(
                                            "  Various options (more than 50 options) - Display skipped by user.",
                                            style='Normal')
                                        p_grouped_options.paragraph_format.left_indent = Inches(1.0)
                                else:
                                    # Small list: always display all options
                                    p_options_label = document.add_paragraph("Options:", style='Normal')
                                    p_options_label.paragraph_format.left_indent = Inches(0.5)

                                    for option in options_for_field:
                                        # Use safe_text_for_docx here
                                        safe_option = safe_text_for_docx(option)
                                        p_option_bullet = document.add_paragraph(safe_option, style='List Bullet')
                                        p_option_bullet.paragraph_format.left_indent = Inches(1.0)
                            else:
                                p_no_options = document.add_paragraph(
                                    "  (No valid options defined or retrieved for this dropdown)", style='Normal')
                                p_no_options.paragraph_format.left_indent = Inches(0.5)
                        else:
                            p = document.add_paragraph()
                            p.add_run(f'{display_field_desc}: ').bold = True
                            p.add_run(f'[{field_type}; {safe_field_id}] ')
                            p.add_run('')

                        processed_field_ids_for_current_subsection.add(field_id)
                    document.add_paragraph()

            # Save the Word document to an in-memory buffer
            doc_buffer = io.BytesIO()
            document.save(doc_buffer)
            doc_buffer.seek(0)

            # Clean filenames
            clean_form_desc = "".join(x for x in form_desc if x.isalnum() or x.isspace()).strip().replace(" ", "_")
            clean_form_id = "".join(x for x in str(form_id) if x.isalnum() or x.isspace()).strip().replace(" ", "_")
            output_filename = f"{clean_form_id}_{today_date_str}_{clean_form_desc}_form.docx"

            generated_files.append((output_filename, doc_buffer.read()))

        if generated_files:
            st.success(f"✅ Successfully generated {len(generated_files)} Word forms.")
        else:
            st.info("ℹ️ No forms were generated from the Excel file.")

        return generated_files

    except pd.errors.EmptyDataError:
        st_display_error("Error", "The Excel file is empty or has no data for form generation.")
        return None
    except Exception as e:
        st_display_error("An Error Occurred", f"An unexpected error occurred during Word form creation: {e}")
        return None


# ----------------------------------------------------------------------------------
# --- ECCAIRS Mappings and Duplicates Documents (REMOVED as requested) ---
# ----------------------------------------------------------------------------------

# REMOVED: def create_eccairs_mappings_document_st(...)
# REMOVED: def create_eccairs_dropdown_mappings_document_st(...)
# REMOVED: def create_missing_eccairs_mappings_document_st(...)
# REMOVED: def create_potential_duplicate_fields_document_st(...)


# --- Streamlit Main Application Logic (Multi-Phase) ---

def reset_app_state():
    """Clears all processing-related keys from session state."""
    keys_to_delete = ['processed_excel_bytes', 'valid_values', 'config_done',
                      'large_dropdowns', 'generated_files', 'file_processed', 'select_all_large_options']
    for key in keys_to_delete:
        if key in st.session_state:
            del st.session_state[key]

    # Also clear all dynamic configuration keys
    dynamic_keys = [k for k in st.session_state.keys() if k.startswith('config_choice_')]
    for key in dynamic_keys:
        del st.session_state[key]


def toggle_all_options():
    """Callback function to set all large option choices based on the Select All state."""
    # The 'select_all_large_options' value in st.session_state is the new value after the toggle.
    select_all = st.session_state['select_all_large_options']

    if 'large_dropdowns' in st.session_state and st.session_state['large_dropdowns']:
        for field_id, _, _ in st.session_state['large_dropdowns']:
            key = f'config_choice_{field_id}'
            # Programmatically set the state for the individual checkboxes
            st.session_state[key] = select_all


def main_app():
    """The main Streamlit application function with multi-step logic."""

    # 💡 Inject custom CSS for the green and red buttons
    st.markdown("""
    <style>
    /* ---------------------------------------------------------------------- */
    /* CSS for the GREEN button (Confirm Configuration and Generate Documents) */
    /* It uses type="secondary" for targeting */
    /* ---------------------------------------------------------------------- */
    div.stButton button[kind="secondary"] {
        background-color: #4CAF50; /* Green background */
        color: white; /* White text */
        border-color: #4CAF50;
    }
    div.stButton button[kind="secondary"]:hover {
        background-color: #45a049 !important; /* Slightly darker green on hover */
        border-color: #45a049 !important;
        color: white !important;
    }
    div.stButton button[kind="secondary"]:focus {
        box-shadow: 0 0 0 0.2rem rgba(76, 175, 80, 0.5); /* Green focus ring */
    }

    /* ---------------------------------------------------------------------- */
    /* CSS for the RED button (Start Over with a New File) */
    /* We target the button explicitly using its Streamlit key (btn_start_over) */
    /* This specificity is required to override the default/green styling. */
    /* ---------------------------------------------------------------------- */
    div[data-testid="stButton-btn_start_over"] > button {
        background-color: #D32F2F !important; /* Red background */
        color: white !important; /* White text */
        border-color: #D32F2F !important;
    }
    div[data-testid="stButton-btn_start_over"] > button:hover {
        background-color: #B71C1C !important; /* Darker red on hover */
        border-color: #B71C1C !important;
        color: white !important;
    }
    div[data-testid="stButton-btn_start_over"] > button:focus {
        box-shadow: 0 0 0 0.2rem rgba(211, 47, 47, 0.5) !important; /* Red focus ring */
        border-color: #D32F2F !important;
    }
    </style>
    """, unsafe_allow_html=True)

    st.title("📝 Excel Export to Word Form Generator")
    st.markdown("Ruben Inion")
    st.markdown(
        "Upload your iQSMS Form/Field export (Excel) to generate Word document forms.") # Adjusted description

    # 1. User Input for Output Folder Name
    folder_name = st.text_input(
        "Enter the base name for your output documents/zip file (e.g. Customer code):",
        st.session_state.get('folder_name', "IQSMS_Forms_Export"),
        key='input_folder_name',
        help="This name will be used to prefix the generated Word files."
    )
    if not folder_name:
        st.stop()
    st.session_state['folder_name'] = folder_name

    # 2. File Uploader
    uploaded_file = st.file_uploader(
        "Choose a form/field Excel file (.xlsx or .xls)",
        type=['xlsx', 'xls'],
        accept_multiple_files=False,
        key='file_uploader'
    )

    st.markdown("---")

    # PHASE 1: Upload and Pre-process
    # Check if a file is uploaded AND if the processed data is NOT yet in state
    if uploaded_file is not None and st.session_state.get('processed_excel_bytes') is None:
        if st.button("1. Start Pre-processing", key='btn_preprocess'):
            # Reset state for a new file/process start
            reset_app_state()

            with st.spinner("Processing Excel and identifying fields..."):
                # Pass the uploaded file object directly, process_excel_st reads the bytes
                processed_excel_bytes, valid_values = process_excel_st(uploaded_file)

                if processed_excel_bytes is not None:
                    st.session_state['processed_excel_bytes'] = processed_excel_bytes
                    st.session_state['valid_values'] = valid_values

                    # Identify large dropdowns here
                    try:
                        # Use the processed bytes to analyze the data
                        temp_df = pd.read_excel(io.BytesIO(processed_excel_bytes))
                        large_dropdowns = []

                        # Use the same column names as used in the generation functions
                        unique_dropdowns = temp_df[
                            temp_df["Field Type"].astype(str).str.strip().str.lower() == "dropdown select"
                            ].drop_duplicates(subset=["Field ID"])[["Field ID", "Field Description"]]

                        for index, row in unique_dropdowns.iterrows():
                            field_id = str(row["Field ID"])
                            field_desc = row["Field Description"]

                            # Count the unique options for this field ID
                            options_count = temp_df[
                                (temp_df["Field ID"].astype(str) == field_id) &
                                (temp_df["Option"].astype(str).str.lower() != "n/a") &
                                (temp_df["Option"].notna())
                                ]["Option"].nunique()

                            if options_count > 50:
                                large_dropdowns.append((field_id, field_desc, options_count))

                        st.session_state['large_dropdowns'] = large_dropdowns
                        st.success("Pre-processing complete. Proceed to Step 2 to configure form generation.")
                        st.session_state['file_processed'] = True
                        st.rerun()  # Rerun to move to Phase 2
                    except Exception as e:
                        st_display_error("Data Analysis Error",
                                         f"Failed to analyze processed data for large dropdowns: {e}")
                        # Clear the processed data to force re-upload
                        st.session_state['processed_excel_bytes'] = None

    # PHASE 2: Configure Large Options (The Pause)
    if st.session_state.get('processed_excel_bytes') is not None and not st.session_state.get('config_done', False):

        large_dropdowns = st.session_state['large_dropdowns']

        st.header("2. Configure Large Option Lists ⚠️")

        if not large_dropdowns:
            st.info("No large dropdown lists (over 50 options) were found. Proceeding to final step.")
            st.session_state['config_done'] = True
            st.rerun()  # Rerun to move to Phase 3
        else:
            st.warning(
                "Please check the boxes below if you want to include all options for fields with **more than 50 options**. Unchecked options will show a simplified entry.")

            with st.container(border=True):
                # --- Select All Checkbox ---
                st.checkbox(
                    "**SELECT ALL** (Display all options for every large dropdown)",
                    value=st.session_state.get('select_all_large_options', False),
                    key='select_all_large_options',
                    on_change=toggle_all_options
                )
                st.markdown("---")

                # Draw the individual checkboxes
                for field_id, field_desc, count in large_dropdowns:
                    key = f'config_choice_{field_id}'

                    # 1. Initialize the choice if not present (this happens on the first run of Phase 2)
                    if key not in st.session_state:
                        # Use the 'select all' state as the default initial value
                        st.session_state[key] = st.session_state.get('select_all_large_options', False)

                    # 2. Draw the checkbox
                    # Explicitly set the value from state
                    st.checkbox(
                        f"Display all **{count}** options for **'{field_desc}'** (ID: {field_id})?",
                        value=st.session_state[key],
                        key=key
                    )

            # NOTE: The custom CSS for the green button is applied via type="secondary"
            if st.button("2. Confirm Configuration and Generate Documents", type="secondary"):
                st.session_state['config_done'] = True
                st.rerun()  # Rerun to move to the next step

    # PHASE 3: Generate and Download
    if st.session_state.get('config_done', False) and st.session_state.get('processed_excel_bytes') is not None:

        st.header("3. Generate Documents & Download")

        # Run generation logic only once per configuration
        if st.session_state.get('generated_files') is None:

            with st.spinner("Generating Word Forms..."):
                # --- Generation ---

                # 1. Main Form Generation
                generated_forms = create_forms_from_excel_st(
                    st.session_state['processed_excel_bytes'],
                    st.session_state['folder_name']
                )

                # All other document generations (ECCAIRS, Duplicates) have been removed.

                # Consolidate results
                all_generated_files = generated_forms if generated_forms else []

                st.session_state['generated_files'] = all_generated_files
                st.success("Document generation complete! Use the buttons below to download.")

        # --- Download Logic (Always runs if files are in state) ---
        if st.session_state.get('generated_files'):
            st.markdown("---")
            st.subheader("Download Generated Files ⬇️")

            all_files_to_zip = st.session_state['generated_files']

            # Individual Downloads
            st.markdown("##### Individual Documents:")
            for filename, file_bytes in all_files_to_zip:
                label_prefix = "Form" # Only Forms are left
                st.download_button(
                    label=f"⬇️ Download {label_prefix}: {filename}",
                    data=file_bytes,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"download_{filename}"
                )

            # Zip Download
            if all_files_to_zip:
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED, False) as zip_file:
                    for filename, file_bytes in all_files_to_zip:
                        zip_file.writestr(filename, file_bytes)

                zip_buffer.seek(0)

                clean_folder_name = "".join(
                    x for x in st.session_state['folder_name'] if x.isalnum() or x.isspace()).strip().replace(" ", "_")
                zip_filename = f"{clean_folder_name}_Generated_Forms.zip" # Updated zip filename

                st.markdown("---")
                st.subheader("All Files in One Zip")
                st.download_button(
                    label=f"📦 Download All {len(all_files_to_zip)} Files as ZIP: {zip_filename}",
                    data=zip_buffer.getvalue(),
                    file_name=zip_filename,
                    mime="application/zip",
                    key="download_all_zip"
                )

        # --- START OVER BUTTON (MOVED HERE) ---
        st.markdown("---")
        # This button is now correctly placed in Phase 3 and styled red by the CSS block.
        if st.button("🔄 Start Over with a New File", key="btn_start_over"):
            reset_app_state()
            st.rerun()


if __name__ == "__main__":
    # Initialize session state variables for multi-step persistence
    if 'processed_excel_bytes' not in st.session_state:
        st.session_state['processed_excel_bytes'] = None
    if 'valid_values' not in st.session_state:
        st.session_state['valid_values'] = set()
    if 'folder_name' not in st.session_state:
        st.session_state['folder_name'] = "IQSMS_Forms_Export"

    # State variables for multi-step process
    if 'config_done' not in st.session_state:
        st.session_state['config_done'] = False
    if 'large_dropdowns' not in st.session_state:
        st.session_state['large_dropdowns'] = []
    if 'generated_files' not in st.session_state:
        st.session_state['generated_files'] = None
    if 'file_processed' not in st.session_state:
        st.session_state['file_processed'] = False

    # State variable for the "Select All" feature
    if 'select_all_large_options' not in st.session_state:
        st.session_state['select_all_large_options'] = False

    main_app()